from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from mm_qbank.io.docx_out import HANDOUT_FONT_NAME, write_lecture_handout_docx


def test_write_lecture_handout_docx_roundtrip(tmp_path: Path) -> None:
    p = tmp_path / "h.docx"
    write_lecture_handout_docx(
        p,
        [
            {
                "题号": "1",
                "修正后问题": "题干A",
                "原问题": "",
                "修正后解析": "解析A",
                "原解析": "",
                "要点": "· 要点1",
                "讲课内容": "口播…",
            }
        ],
    )
    assert p.is_file()
    assert p.stat().st_size > 800


def test_docx_renders_markdown_double_star_as_word_bold(tmp_path: Path) -> None:
    p = tmp_path / "bold.docx"
    write_lecture_handout_docx(
        p,
        [
            {
                "题号": "1",
                "修正后问题": "关于**中心静脉压（CVP）**监测的叙述",
                "原问题": "",
                "修正后解析": "解析",
                "原解析": "",
                "要点": "",
                "讲课内容": "",
            }
        ],
    )
    d = Document(p)
    para_with_cvp = next(
        (para for para in d.paragraphs if "中心静脉压" in para.text and "CVP" in para.text),
        None,
    )
    assert para_with_cvp is not None
    assert "*" not in para_with_cvp.text
    assert any(run.bold for run in para_with_cvp.runs)


def test_points_lines_become_word_list_bullets(tmp_path: Path) -> None:
    p = tmp_path / "bullets.docx"
    pts = (
        "· 第一点\n"
        "· 第二点含**加粗**\n"
        "- Markdown 式一行\n"
        "非列表行单独一段"
    )
    write_lecture_handout_docx(
        p,
        [
            {
                "题号": "1",
                "修正后问题": "题",
                "原问题": "",
                "修正后解析": "",
                "原解析": "",
                "要点": pts,
                "讲课内容": "",
            }
        ],
    )
    d = Document(p)
    listed = [
        para
        for para in d.paragraphs
        if para.style and getattr(para.style, "name", "") == "List Bullet"
    ]
    assert len(listed) >= 3
    texts = [para.text.replace("\u200b", "") for para in listed]
    assert any("第一点" in t for t in texts)
    assert any("第二点" in t and "加粗" in t for t in texts)
    assert any("Markdown" in t for t in texts)


def test_docx_sets_eastasia_rfonts(tmp_path: Path) -> None:
    p = tmp_path / "ea.docx"
    write_lecture_handout_docx(
        p,
        [
            {
                "题号": "1",
                "修正后问题": "中文测试",
                "原问题": "",
                "修正后解析": "",
                "原解析": "",
                "要点": "",
                "讲课内容": "",
            }
        ],
    )
    with zipfile.ZipFile(p) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "w:eastAsia" in xml
    assert HANDOUT_FONT_NAME in xml
