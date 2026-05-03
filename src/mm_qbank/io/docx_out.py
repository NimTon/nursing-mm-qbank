from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from mm_qbank.io.result_row_qa import handout_stem_analysis

# 讲义 Word 全文字体（Windows 随附 UI 版雅黑）
HANDOUT_FONT_NAME = "Microsoft YaHei UI"


def _apply_handout_font_to_run(run: Any, *, pt: Any | None = None) -> None:
    """
    Word 对汉字实际读 ``w:rFonts/@w:eastAsia``；仅设 ``run.font.name`` 往往只写 ascii/hAnsi，
    界面仍显示宋体等。此处四类脚本统一为同一字体。
    """
    from docx.oxml.ns import qn

    fn = HANDOUT_FONT_NAME
    run.font.name = fn
    if pt is not None:
        run.font.size = pt
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), fn)
    rfonts.set(qn("w:cs"), fn)


def _apply_handout_font_to_paragraph_style(style: Any, *, pt: Any | None = None) -> None:
    """段落样式（Normal / Heading / List Bullet 等）上的默认字体 + 东亚字体。"""
    from docx.oxml.ns import qn

    fn = HANDOUT_FONT_NAME
    style.font.name = fn
    if pt is not None:
        style.font.size = pt
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), fn)
    rfonts.set(qn("w:cs"), fn)


# 行首「点标记」：间隔号 ·、项目符号 •，以及 Markdown 式 ``- `` / ``* ``（须有空格）
_RE_DOT_BULLET = re.compile(r"^\s*[·•‧]\s*(.*)$")
_RE_MD_BULLET = re.compile(r"^\s*[-*＊]\s+(.+)$")


def _md_merge_star_parts(text: str) -> list[str]:
    parts = (text or "").split("**")
    if len(parts) % 2 == 0:
        parts[-2] = parts[-2] + "**" + parts[-1]
        parts = parts[:-1]
    return parts


def _add_runs_md_bold_to_paragraph(p: Any, text: str, *, pt: Any) -> None:
    """向已有段落追加 run：识别 ``**加粗**``。"""
    raw = text or ""
    if not raw.strip():
        r = p.add_run("（无）")
        _apply_handout_font_to_run(r, pt=pt)
        return
    for i, seg in enumerate(_md_merge_star_parts(raw)):
        if seg == "":
            continue
        run = p.add_run(seg)
        run.bold = i % 2 == 1
        _apply_handout_font_to_run(run, pt=pt)


def _bullet_line_payload(line: str) -> str | None:
    """
    若本行为「点标记」列表行，返回去掉标记后的正文；否则返回 None。
    支持：· • ‧ 以及 ``- `` / ``* ``（须后跟空格，避免误伤「5–12」类连字符）。
    """
    s = line.strip()
    if not s:
        return None
    m = _RE_DOT_BULLET.match(s)
    if m:
        return m.group(1).strip()
    m = _RE_MD_BULLET.match(s)
    if m:
        return m.group(1).strip()
    return None


def _append_paragraph_markdown_double_star_bold(doc: Any, text: str, *, pt: Any) -> None:
    """单段正文：``**加粗**``。"""
    p = doc.add_paragraph()
    raw = text or ""
    if not raw.strip():
        r = p.add_run("（无）")
        _apply_handout_font_to_run(r, pt=pt)
        return
    _add_runs_md_bold_to_paragraph(p, raw, pt=pt)


def _append_points_with_native_bullets(doc: Any, text: str, *, pt: Any) -> None:
    """
    【要点】专用：行首 ``·`` / ``•`` / ``- `` / ``* `` 转为 Word 原生项目符号段，
    段内仍支持 ``**加粗**``。
    """
    raw = (text or "").strip()
    if not raw:
        _append_paragraph_markdown_double_star_bold(doc, "", pt=pt)
        return

    lines = raw.splitlines()
    list_bullet_style = "List Bullet"
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        payload = _bullet_line_payload(s)
        if payload is not None:
            try:
                p = doc.add_paragraph(style=list_bullet_style)
            except (KeyError, ValueError):
                p = doc.add_paragraph()
                _add_runs_md_bold_to_paragraph(p, "· " + payload, pt=pt)
                continue
            _add_runs_md_bold_to_paragraph(p, payload, pt=pt)
        else:
            p = doc.add_paragraph()
            _add_runs_md_bold_to_paragraph(p, s, pt=pt)


def write_lecture_handout_docx(path: Path, rows: list[dict[str, Any]]) -> None:
    """
    按题生成讲义 Word：每题含【题目】【解析】【要点】【讲课内容】；
    题目与解析为修正后定稿（无修正列时回退原列）。
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt  # type: ignore[import-untyped]
    except ModuleNotFoundError as e:
        raise ModuleNotFoundError(
            "导出 Word 需要安装 python-docx：pip install python-docx"
        ) from e

    path = path.resolve()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_handout_font_to_paragraph_style(doc.styles["Normal"], pt=Pt(11))
    for _sid in ("List Bullet", "List Paragraph"):
        try:
            _apply_handout_font_to_paragraph_style(doc.styles[_sid], pt=Pt(11))
        except KeyError:
            pass
    try:
        _apply_handout_font_to_paragraph_style(doc.styles["Heading 1"], pt=Pt(14))
    except KeyError:
        pass

    for i, row in enumerate(rows):
        tihao = str(row.get("题号") or "").strip()
        stem, analysis = handout_stem_analysis(row)
        points = str(row.get("要点") or "").strip()
        lecture = str(row.get("讲课内容") or "").strip()

        if i > 0:
            doc.add_page_break()

        title = f"第 {tihao} 题" if tihao else f"第 {i + 1} 题"
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            _apply_handout_font_to_run(run, pt=Pt(14))

        qt = str(row.get("题目类型") or "").strip()
        if qt:
            pq = doc.add_paragraph()
            rq = pq.add_run(f"【题目类型】{qt}")
            rq.bold = True
            _apply_handout_font_to_run(rq, pt=Pt(11))

        def _section(heading: str, body: str, *, points_block: bool = False) -> None:
            p0 = doc.add_paragraph()
            r0 = p0.add_run(heading)
            r0.bold = True
            _apply_handout_font_to_run(r0, pt=Pt(11))
            if points_block:
                _append_points_with_native_bullets(doc, body, pt=Pt(11))
            else:
                _append_paragraph_markdown_double_star_bold(doc, body, pt=Pt(11))

        _section("【题目】", stem)
        _section("【解析】", analysis)
        _section("【要点】", points, points_block=True)
        _section("【讲课内容】", lecture)

    doc.save(path)
